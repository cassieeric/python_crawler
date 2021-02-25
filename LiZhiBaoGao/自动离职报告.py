
# -*- coding: utf-8 -*-
"""
@Time ：2021/2/1 10:51
@Auth ：wutong
@File ：自动离职报告.py
@IDE ：PyCharm
"""
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import RGBColor, Pt
from docx.shared import Inches, RGBColor
from lzStr import lzStr

# todo 初始化一个文档
document = Document()
# 全局指定字体
document.styles['Normal'].font.name = u'.萍方-简'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'.萍方-简')

# todo 加个标题
paragraph = document.add_heading('离职申请', level=3)

# todo 居中
paragraph_format = paragraph.paragraph_format
paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

paragraph = document.add_paragraph()
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = paragraph.add_run("")
run.add_picture('./image.jpg', width=Inches(1.0), height=Inches(1.0))

paragraph = document.add_paragraph()
lz_add_run = paragraph.add_run(lzStr.format(name="吴老板"))
lz_add_run.font.size = Pt(8)  # 字体大小设置，和word里面的字号相对应
lz_add_run.font.color.rgb = RGBColor(54, 95, 145)

document.save('离职.docx')